from PyQt5 import uic
from PyQt5.QtWidgets import QApplication, QMainWindow, QTableWidgetItem, QFrame, QPushButton, QLabel, QVBoxLayout, QLineEdit, QHBoxLayout, QComboBox, QToolButton, QMessageBox, QFileDialog, QProgressBar, QHeaderView, QSizePolicy, QDialog, QWidget, QGraphicsBlurEffect
from PyQt5.QtGui import QIcon, QCursor, QFont, QMovie
from PyQt5.QtCore import Qt, QThread, pyqtSignal
import rssrce
import pandas as pd
import numpy as np
import sqlite3
import math
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
import os
import sys
import json
import ctypes
import logging
from datetime import date, timedelta
from cryptography.fernet import Fernet
from pathlib import Path

SECRET_KEY = b'bdhRPNACdRgoFcAVIDItlKH_ffYkAu8i4I6V0-nXf5k='

# Chemin AppData/Roaming/Sige
appdata = os.getenv("APPDATA")
folder = os.path.join(appdata, "Sige")
os.makedirs(folder, exist_ok=True)
license_path = os.path.join(folder, "licence.bin")
docSaveDir = str(Path.home() / "Documents") # Chemin de départ pour le QFileDialog

def resource_path(relative_path):
    """ Trouve le chemin du fichier UI, que ce soit en mode script ou exécutable """
    if getattr(sys, 'frozen', False):  # Si PyInstaller exécute l'EXE
        base_path = sys._MEIPASS
    else:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

# Configurer le logger pour écrire dans un fichier
logging.basicConfig(filename=os.path.join(folder, "log.txt"), level=logging.ERROR, format="%(asctime)s - %(levelname)s - %(message)s")

def write_licence(expiration_date: date):

    data = {
        "expires": expiration_date.isoformat()  # YYYY-MM-DD
    }

    json_bytes = json.dumps(data).encode("utf-8")

    fernet = Fernet(SECRET_KEY)
    encrypted = fernet.encrypt(json_bytes)

    with open(license_path, "wb") as f:
        f.write(encrypted)

    # cacher le fichier
    FILE_ATTRIBUTE_HIDDEN = 0x02
    ctypes.windll.kernel32.SetFileAttributesW(license_path, FILE_ATTRIBUTE_HIDDEN)

def read_licence() -> bool:

    if not os.path.exists(license_path):
        return False

    try:
        with open(license_path, "rb") as f:
            encrypted = f.read()

        fernet = Fernet(SECRET_KEY)
        decrypted = fernet.decrypt(encrypted)

        data = json.loads(decrypted.decode("utf-8"))
        expires = date.fromisoformat(data["expires"])

        return date.today() <= expires

    except Exception:
        return False

# Charger l'interface .ui
class MyApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowIcon(QIcon(resource_path("icons/app_icon.png")))
        self.setWindowTitle("Sige-Decisions")
        self.setMinimumSize(900,600)
        con = sqlite3.connect(os.path.join(folder, "database.db"))

        try:
            self.data = pd.read_sql_query("select * from academy",con)
            self.data.sort_values(by=["AE","Centre"],inplace=True)
        except:
            cur = con.cursor()
            cur.execute("""
                CREATE TABLE IF NOT EXISTS "academy" (
                "AE"	TEXT,
                "Centre"	TEXT,
                "Responsabilité"	TEXT,
                "Prénom"	TEXT,
                "Nom"	TEXT,
                "Matricule"	TEXT,
                "Service"	TEXT,
                "Catégorie"	TEXT,
                "Poste"	TEXT,
                "Telephone"	TEXT,
                "Examen"	TEXT,
                "Nb salle"	REAL
            );
            """)
            con.commit()
            self.data = pd.read_sql_query("select * from academy",con)
            self.data.sort_values(by=["AE","Centre"],inplace=True)
            con.close()
        fichier_ui = resource_path("main.ui")
        self.widget = uic.loadUi(fichier_ui)  # Charger le fichier .ui
        self.statusBar().setStyleSheet("""
            background-color: qlineargradient(
                spread:pad,
                x1:0.481, y1:0.0397727,
                x2:0.497373, y2:0.517,
                stop:0.109 rgba(10, 21, 72, 255),
                stop:1 rgba(24, 29, 52, 255)
            );
        """)
        self.widget.table.setSizePolicy(
            QSizePolicy.Expanding,
            QSizePolicy.Expanding
        )
        self.widget.table.horizontalHeader().setStretchLastSection(True)
        self.widget.table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.widget.add_button.clicked.connect(lambda: self.add_slot(self.data.AE.unique()))
        self.widget.generer.clicked.connect(self.generer_word)
        self.widget.upload_btn.clicked.connect(self.upload)
        self.widget.edit.clicked.connect(self.modify)
        self.widget.reset_btn.clicked.connect(self.reset_slot)
        self.widget.export_2.clicked.connect(self.export_slot)
        self.setup()

        # La page à afficher si la date est expirée
        unavailable = QFrame()
        unavailable.setStyleSheet("background-color: qlineargradient(spread:pad, x1:0.481, y1:0.0397727, x2:0.471511, y2:0.511, stop:0.109 rgba(10, 21, 72, 255), stop:1 rgba(24, 29, 52, 255));")
        unavailable_label = QLabel()
        unavailable_label.setText("Session expirée. Veuillez contacter l'administrateur.")
        unavailable_label.setStyleSheet("color: red; background: transparent; border: none")
        unavailable_label.setFont(QFont("Arial Black", 18, QFont.Bold))
        unavailable_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        unavailable_layout = QVBoxLayout()
        unavailable_layout.addWidget(unavailable_label)
        unavailable.setLayout(unavailable_layout)

        start = date(2025, 12, 1)
        expiration = start + timedelta(days=217)

        # 1️ créer la licence UNE SEULE FOIS
        if not os.path.exists(license_path):
            write_licence(expiration)

        print(read_licence())
        # 2️ vérifier la licence
        if read_licence():
            self.setCentralWidget(self.widget)
        else:
            self.setCentralWidget(unavailable)

    def setup(self):
        academies = self.data['AE'].unique().tolist()
        self.widget.academy_combo.addItems(academies)
        self.widget.academy_combo.currentIndexChanged.connect(self.combo_change)

        self.widget.searchLine.textChanged.connect(self.search_slot)
        
        self.widget.table.setRowCount(self.data.shape[0])
        self.widget.table.setColumnCount(self.data.shape[1])

        self.widget.table.setHorizontalHeaderLabels(self.data.columns)

        for row in range(self.data.shape[0]):
            for col in range(self.data.shape[1]):
                item = QTableWidgetItem(str(self.data.iat[row, col]))
                font = QFont()
                font.setPointSize(9)  # taille de police souhaitée
                item.setFont(font)
                self.widget.table.setItem(row, col, item)  # Remplissage des cellules
        self.statusBarProgress = QProgressBar()
        self.statusBarProgress.setStyleSheet("""
            QProgressBar {
                background-color: rgba(255, 255, 255, 0.12);
                border-radius: 10px;
                text-align: center;
                color: #eaeaea;
                height: 14px;
            }

            QProgressBar::chunk {
                background-color: rgba(90, 160, 255, 0.85);
                border-radius: 10px;
            }
        """)
        self.statusBarProgress.setMinimumWidth(300)
        self.statusBarProgress.setVisible(False)
        self.statusBar().addPermanentWidget(self.statusBarProgress, 2)
        version_label = QLabel("Sige Decision v26")
        version_label.setStyleSheet("background: transparent; color: #fff;")
        self.statusBar().addPermanentWidget(version_label)

    def search_slot(self):
        conn = sqlite3.connect(os.path.join(folder, "database.db"))
        req = """SELECT * FROM academy WHERE Centre LIKE ? OR 
                AE LIKE ? 
                OR Matricule LIKE ? 
                OR Prénom LIKE ? 
                OR Nom LIKE ?
                OR Telephone LIKE ?
                OR Service LIKE ?
                OR Poste LIKE ?
                OR Catégorie LIKE ?
                OR Responsabilité LIKE ?
        """
        param = f"%{self.widget.searchLine.text()}%"
        result = pd.read_sql_query(req, conn, params=(param, param, param, param, param, param, param, param, param, param))
        conn.close()
        self.tableFill(result, self.widget.table)
        

    def fillTableWidget(self, rows, columns, tableWidget):
        tableWidget.setUpdatesEnabled(False)

        tableWidget.setRowCount(len(rows))
        tableWidget.setColumnCount(len(columns))
        tableWidget.setHorizontalHeaderLabels(columns)

        font = QFont()
        font.setPointSize(9)

        for r, row in enumerate(rows):
            for c, value in enumerate(row):
                item = QTableWidgetItem(value)
                item.setFont(font)
                tableWidget.setItem(r, c, item)

        tableWidget.setUpdatesEnabled(True)

    def tableFill(self, data: pd.DataFrame, tableWidget):
        data.sort_values(by=["AE","Centre"],inplace=True)
        self.statusBarProgress.setValue(0)
        self.statusBarProgress.show()

        self.thread = TableFillThread(data)

        self.thread.progress.connect(self.statusBarProgress.setValue)

        self.thread.finished.connect(
            lambda rows, cols: self.fillTableWidget(rows, cols, tableWidget)
        )

        self.thread.finished.connect(self.thread.quit)
        self.thread.finished.connect(self.statusBarProgress.hide)

        self.thread.error.connect(
            lambda cols: tableWidget.setHorizontalHeaderLabels(cols)
        )

        self.thread.start()

    """
    def tableFill(self, data:pd.DataFrame, tableWidget:None):
        try:
            tableWidget.setRowCount(data.shape[0])
            tableWidget.setColumnCount(data.shape[1])

            tableWidget.setHorizontalHeaderLabels(data.columns)

            for row in range(data.shape[0]):
                for col in range(data.shape[1]):
                    item = QTableWidgetItem(str(data.iat[row, col]))
                    font = QFont()
                    font.setPointSize(9)  # taille de police souhaitée
                    item.setFont(font)
                    tableWidget.setItem(row, col, item)  # Remplissage des cellules
        except:
            columns = ["AE","Centre","Responsabilité","Prénom","Nom","Matricule","Service","Catégorie","Poste","Telephone","Examen","Nb salle"]
            tableWidget.setRowCount(len(columns))
            tableWidget.setHorizontalHeaderLabels(columns)
    """
        
    def combo_change(self):
        academie_selected = self.widget.academy_combo.currentText()
        con = sqlite3.connect(os.path.join(folder, "database.db"))
        if academie_selected != "Tout":
            try:
                self.data = pd.read_sql_query(f"select * from academy WHERE AE='{academie_selected}'",con)
                self.data.sort_values(by=["AE","Centre"],inplace=True)
            except:
                QMessageBox.critical(self, "Erreur", f"L'academie {academie_selected} n'est pas dans la base de données !")
            finally:
                self.tableFill(self.data, self.widget.table)
        else:
            try:
                self.data = pd.read_sql_query("select * from academy",con)
                self.data.sort_values(by=["AE","Centre"],inplace=True)
            except:
                QMessageBox.critical(self, "Erreur", "La base de données est vide !")
            finally:
                self.tableFill(self.data, self.widget.table)
        con.close()
    def upload(self):
        options = QFileDialog.Options()
        options |= QFileDialog.DontUseNativeDialog
        file_path, _ = QFileDialog.getOpenFileName(self, "Choisir un fichier", docSaveDir if docSaveDir else "", "Tous les fichiers (*);;Fichiers Excel (*.xlsx)", options=options)
        if file_path:
            try:
                loading = LoadingPage()
                loading.show()
                thread = DataUploadThread(file_path)
                thread.academieList.connect(lambda academies_list: self.widget.academy_combo.addItems(academies_list))
                thread.concatenateDf.connect(lambda concatened_df: self.tableFill(concatened_df, self.widget.table))
                thread.finished.connect(lambda: loading.close())
                thread.finished.connect(lambda success: QMessageBox.information(self, "Succès", "Vos données ont été importées avec succès !") if success else QMessageBox.critical(self, "Erreur", "Une erreur est survenue lors de l'importation des données."))
                thread.finished.connect(thread.quit)
                thread.finished.connect(thread.wait)
                thread.start()
            except sqlite3.IntegrityError as e:
                QMessageBox.critical(self, "Erreur", f"Les informations de certaines personnes de cet fichier sont soit dupliquées ou existent déjà dans la base de données.")
                logging.error("Une erreur est survenue", exc_info=True)  # Enregistre l'erreur avec stack trace
    def add_slot(self, academies:list=[]):
        fichier_ui = resource_path("add.ui")

        self.add_widget = uic.loadUi(fichier_ui)
        self.add_widget.academy_combo.clear()
        self.add_widget.academy_combo.addItems(academies)
        self.add_widget.academy_add_btn.setIcon(QIcon(resource_path("icons/plus.svg")))
        self.add_widget.academy_add_btn.clicked.connect(lambda: self.academie_add(academies, self.add_widget.academy_combo))

        dialog = FrostedDialog(self, self.add_widget)
        self.add_widget.add_close.clicked.connect(dialog.close)

        self.vice_group_layout = QVBoxLayout()
        self.vice_group_layout.setContentsMargins(9, 25, 9, 9)
        self.vice_group_layout.addWidget(Vice_president("prénom0","nom0","matricule0","service0","categorie0","poste0","telephone0",self.vice_group_layout, "new"))

        Vice_president_add = QToolButton()
        Vice_president_add.setStyleSheet("background-color: rgba(37, 107, 182, 200); border-radius: 8px; border: 1px solid rgba(255,255,255,0.40)")
        Vice_president_add.setText("Ajouter")
        Vice_president_add.setToolTip("Ajouter un vice président")
        Vice_president_add.setFixedHeight(40)
        Vice_president_add.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
        Vice_president_add.setCursor(QCursor(Qt.CursorShape.PointingHandCursor))
        Vice_president_add.clicked.connect(lambda: self.vice_president_add_slot(self.vice_group_layout))

        self.vice_group_layout.addWidget(Vice_president_add)
        self.add_widget.vice_group.setLayout(self.vice_group_layout)

        self.add_widget.enregistrer.clicked.connect(lambda: self.save_slot(dialog))

        content_layout = dialog.container.layout()
        content_layout.addWidget(self.add_widget)

        dialog.exec()
        
    def academie_add(self, academiesList=None, academieCombo=None):
        win = QDialog(self)
        win.setWindowTitle("Ajouter une academie")
        layout = QVBoxLayout()

        line = QLineEdit()
        line.setPlaceholderText("Nom de l'academie")
        layout.addWidget(line)

        ok = QPushButton("Ajouter")
        ok.clicked.connect(lambda: self.academie_add_save(academiesList, academieCombo, line, win))
        cancel = QPushButton("Annuler")
        cancel.clicked.connect(win.close)
        btnsLayout = QHBoxLayout()
        btnsLayout.addWidget(ok)
        btnsLayout.addWidget(cancel)
        layout.addLayout(btnsLayout)
        win.setLayout(layout)

        win.exec()
    def academie_add_save(self, academiesList:list, academieCombo, lineEdit, dialog):
        """Fonction pour valider l'ajout d'une nouvelle academie.
        Params:
        - academiesList: list (liste des academies existantes)
        - academieCombo: QComboBox (comboBox des academies de la boîte d'ajout)
        - lineEdit: QLineEdit (champ de texte pour le nom de la nouvelle academie)
        - dialog: QDialog (boîte de dialogue pour ajouter une academie)
        """
        new_academie = lineEdit.text()
        if new_academie == "":
            msg = QMessageBox()
            msg.setWindowTitle("Erreur")
            msg.setText("Veuillez entrer le nom d'une academie")
            msg.setIcon(QMessageBox.Icon.Critical)
            msg.exec()
        elif new_academie in academiesList:
            msg = QMessageBox()
            msg.setWindowTitle("Erreur")
            msg.setText("Cette academie existe déjà")
            msg.setIcon(QMessageBox.Icon.Critical)
            msg.exec()
        else:
            academieCombo.addItem(new_academie)
            dialog.close()
    def vice_president_add_slot(self, layout:QVBoxLayout):
        layout_children = layout.count()
        c = layout_children - 1 # Enlever 1 pour ne pas compter le button d'ajout de vice président
        arguments = [f"prénom{c}",f"nom{c}",f"matricule{c}",f"service{c}",f"categorie{c}",f"poste",f"telephone{c}"]
        layout.insertWidget(c,Vice_president(arguments[0],arguments[1],arguments[2],arguments[3],arguments[4],
                                                        arguments[5],arguments[6],layout,status="new"))

    # Fonction qui collecte les informations et les insérées dans la base de données
    def save_slot(self, dialog=None):
        academy = self.add_widget.academy_combo.currentText()
        centre = self.add_widget.centre_line.text()
        examen = self.add_widget.examen_combo.currentText()
        nb_salle = self.add_widget.salle_spin.value()
        msg = QMessageBox()
        msg.setWindowTitle("Erreur")
        if academy == "Academie d'enseignement":
            msg.setText("Veuillez selectionner une academie valide")
            msg.setIcon(QMessageBox.Icon.Critical)
            msg.exec()
        elif centre == "":
            msg.setText("Veuillez ecrire un centre d'examen")
            msg.setIcon(QMessageBox.Icon.Critical)
            msg.exec()
        elif examen == "Examen":
            msg.setText("Veuillez selectionner un examen valide")
            msg.setIcon(QMessageBox.Icon.Critical)
            msg.exec()
        elif nb_salle <= 0:
            msg.setText("Le nombre de salle ne peut pas être 0")
            msg.setIcon(QMessageBox.Icon.Critical)
            msg.exec()
        else:
            #Info du président
            presi_prenom = self.add_widget.presi_prenom.text()
            presi_nom = self.add_widget.presi_nom.text()
            presi_matricule = self.add_widget.presi_matricule.text()
            presi_service = self.add_widget.presi_service.text()
            presi_categorie = self.add_widget.presi_categorie.currentText()
            presi_poste = self.add_widget.presi_poste.text()
            presi_telephone = self.add_widget.presi_telephone.text()
            #Construction de DataFrame avec les infos de base et celles du président
            data = {
                "AE": [academy], 
                "Centre": [centre], 
                "Responsabilité": ["Président"], 
                "Prénom": [presi_prenom], 
                "Nom": [presi_nom],
                "Matricule": [presi_matricule], 
                "Service": [presi_service], 
                "Catégorie": [presi_categorie], 
                "Poste": [presi_poste], 
                "Telephone": [presi_telephone],
                "Examen": [examen], 
                "Nb salle": [nb_salle]
            }
            df = pd.DataFrame(data)
            
            vice_presi_count = self.vice_group_layout.count() - 1 # Avoir le nombre de vice présidents

            if nb_salle/3 > vice_presi_count:
                msg.setText("Votre enregistrement ne respecte pas la politique de 3 salles par vice président.\nVoulez-vous continuer ?")
                msg.setIcon(QMessageBox.Icon.Question)
                msg.setStandardButtons(QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
                reponse = msg.exec()
                if reponse == QMessageBox.StandardButton.Yes:
                    for i in range(vice_presi_count):
                        widget = self.vice_group_layout.itemAt(i).widget() # Récupérer le widget de chaque vice président
                        prenom = widget.prenom_line.text()
                        nom = widget.nom_line.text()
                        matricule = widget.matricule_line.text()
                        service = widget.service_line.text()
                        categorie = widget.categorie_combo.currentText()
                        poste = widget.poste_line.text()
                        telephone = widget.telephone_line.text()
                        data = {
                            "AE": academy, 
                            "Centre": centre, 
                            "Responsabilité": f"Vice Président {i+1}", 
                            "Prénom": prenom, 
                            "Nom": nom,
                            "Matricule": matricule, 
                            "Service": service, 
                            "Catégorie": categorie, 
                            "Poste": poste, 
                            "Telephone": telephone,
                            "Examen": examen, 
                            "Nb salle": nb_salle
                        }
                        df.loc[len(df)] = data # Ajouter une ligne au DataFrame avec les infos de base et celle du vice président parcouru
                    self.data = pd.concat([self.data, df], ignore_index=True) # Concatener la DataFrama principale et celle créer après insertion
                    # Réconstruire le tableau d'affichage
                    self.tableFill(self.data, self.widget.table)
                    if dialog:
                        dialog.close()
                    with sqlite3.connect(os.path.join(folder, "database.db")) as conn:
                        self.data.to_sql(name="academy",con=conn, if_exists="replace",index=False)
                    QMessageBox.information(self, "Succès", "Ajout effectué avec succès !")
            else:
                for i in range(vice_presi_count):
                    widget = self.vice_group_layout.itemAt(i).widget() # Récupérer le widget de chaque vice président
                    prenom = widget.prenom_line.text()
                    nom = widget.nom_line.text()
                    matricule = widget.matricule_line.text()
                    service = widget.service_line.text()
                    categorie = widget.categorie_combo.currentText()
                    poste = widget.poste_line.text()
                    telephone = widget.telephone_line.text()
                    data = {
                        "AE": academy, 
                        "Centre": centre, 
                        "Responsabilité": f"Vice Président {i+1}", 
                        "Prénom": prenom, 
                        "Nom": nom,
                        "Matricule": matricule, 
                        "Service": service, 
                        "Catégorie": categorie, 
                        "Poste": poste, 
                        "Telephone": telephone,
                        "Examen": examen, 
                        "Nb salle": nb_salle
                    }
                    df.loc[len(df)] = data # Ajouter une ligne au DataFrame avec les infos de base et celle du vice président parcouru
                self.data = pd.concat([self.data, df], ignore_index=True) # Concatener la DataFrama principale et celle créer après insertion
                # Réconstruire le tableau d'affichage
                self.tableFill(self.data, self.widget.table)
                if dialog:
                    dialog.close()
                with sqlite3.connect(os.path.join(folder, "database.db")) as conn:
                    self.data.to_sql(name="academy",con=conn,if_exists="replace",index=False)
                    self.widget.academy_combo.addItem(data.get('AE'))
                QMessageBox.information(self, "Succès", "Ajout effectué avec succès !")
    def reset_slot(self):
        try:
            msg = QMessageBox()
            msg.setText("Voulez-vous vraiment supprimer tout ?")
            msg.setIcon(QMessageBox.Icon.Question)
            msg.setStandardButtons(QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
            reponse = msg.exec()
            if reponse == QMessageBox.StandardButton.Yes:
                conn = sqlite3.connect(os.path.join(folder, "database.db"))
                cur = conn.cursor()
                cur.execute("DELETE FROM academy")
                cur.execute("""
                    CREATE TABLE IF NOT EXISTS "academy" (
                    "AE"	TEXT,
                    "Centre"	TEXT,
                    "Responsabilité"	TEXT,
                    "Prénom"	TEXT,
                    "Nom"	TEXT,
                    "Matricule"	TEXT,
                    "Service"	TEXT,
                    "Catégorie"	TEXT,
                    "Poste"	TEXT,
                    "Telephone"	TEXT,
                    "Examen"	TEXT,
                    "Nb salle"	REAL
                );
                """)
                conn.commit()
                conn.close()
                self.widget.table.clear()
        except Exception as e:
            logging.error("Une erreur est survenue", exc_info=True)  # Enregistre l'erreur avec stack trace

    def generer_word(self):
        loading_dialog = LoadingPage()
        loading_dialog.show()
        self.statusBarProgress.setVisible(True)
        thread = GenerateWordThread(self)
        thread.progress.connect(lambda value: self.statusBarProgress.setValue(value))
        thread.request_save_path.connect(lambda: self.word_open_save_dialog(thread))
        thread.save_path_received.connect(lambda path: self.save_doc(path, thread))
        thread.finish.connect(lambda: loading_dialog.close())
        thread.finish.connect(lambda success: QMessageBox.information(self, "Succès", "Décision générée avec succès !") if success else QMessageBox.warning(self, "Annulé", "Génération du document annulée."))
        thread.finish.connect(lambda: self.statusBarProgress.setValue(0))
        thread.finish.connect(lambda: self.statusBarProgress.setVisible(False))
        thread.finish.connect(thread.quit)
        thread.finish.connect(thread.wait)

        thread.start()

    def word_open_save_dialog(self, thread:QThread):
        options = QFileDialog.Options()
        options |= QFileDialog.DontUseNativeDialog

        start_dir = str(Path.home() / "Documents")

        path, _ = QFileDialog.getSaveFileName(
            self,
            "Enregistrer le fichier",
            start_dir,
            "Documents Word (*.docx)"
        )

        thread.save_path_received.emit(path)

    def save_doc(self, path, thread):
        if path:
            if not path.endswith(".docx"):
                path += ".docx"
            thread.doc.save(path)
            thread.finish.emit(True)
        else:
            thread.finish.emit(False)

    def export_slot(self):
        try:
            conn = sqlite3.connect(os.path.join(folder, "database.db"))
            df = pd.read_sql_query("SELECT * FROM academy",conn)
            df.sort_values(by=["AE","Centre"],inplace=True)
            options = QFileDialog.Options()
            options |= QFileDialog.DontUseNativeDialog
            file_path, _ = QFileDialog.getSaveFileName(self, "Enregistrer le fichier", docSaveDir if docSaveDir else "", "Tous les fichiers (*)", options=options)

            if file_path:  # Vérifie si un chemin a été sélectionné
                if not file_path.endswith(".xlsx"):  # Ajoute l'extension si nécessaire
                    file_path += ".xlsx"
                df.to_excel(file_path,index=True)
        except Exception as e:
            logging.error("Une erreur est survenue", exc_info=True)  # Enregistre l'erreur avec stack trace
    # Fonction permettant d'afficher la fenetre de modification
    def modify(self):
        deleteVar = [1]  # Liste des matricules à supprimer
        try:
            mod_widget = resource_path("add.ui")
            widget = uic.loadUi(mod_widget)

            widget.academy_add_btn.setVisible(True)

            dialog = FrostedDialog(self, widget)
            widget.add_close.clicked.connect(dialog.close)

            widget.label.setText("Modification")
            widget.presi_matricule.setReadOnly(True)
            # Liste des centre pour le combobox de selection
            self.items = self.data['Centre'].unique().tolist()
            # ComboBox
            self.modify_combo = QComboBox(self)
            self.modify_combo.addItems(self.items)  # Ajout des éléments initiaux
            self.modify_combo.setEditable(True)  # Permet de taper dans la comboBox
            self.modify_combo.currentTextChanged.connect(lambda: self.modify_data(widget, deleteVar))
            self.modify_combo.setStyleSheet(
                """
                QComboBox {
                    background: rgba(255, 255, 255, 0.20);    /* fond translucide pour effet glass */
                    color: #000;                              /* texte sombre ou clair selon ton thème */
                    border: 1px solid rgba(255, 255, 255, 0.25);
                    border-radius: 8px;
                    padding: 6px 30px 6px 10px;              /* espace pour la flèche */
                    font-size: 15px;
                    min-height: 28px;
                }

                /* Hover : éclaircissement léger */
                QComboBox:hover {
                    background: rgba(255, 255, 255, 0.30);
                    border: 1px solid rgba(120, 150, 255, 0.6);
                }

                /* Focus : bordure plus visible */
                QComboBox:focus {
                    background: rgba(255, 255, 255, 0.45);
                    border: 1px solid rgba(120, 150, 255, 0.8);
                }

                /* Flèche de la combobox */
                QComboBox::drop-down {
                    subcontrol-origin: padding;
                    subcontrol-position: top right;
                    width: 25px;
                    border-left: none;
                }

                QComboBox::down-arrow {
                    image: url("icons/arrow-down.svg");  /* flèche personnalisée */
                    width: 12px;
                    height: 12px;
                }

                /* Menu déroulant */
                QComboBox QAbstractItemView {
                    background: rgba(255, 255, 255, 0.25);  /* menu semi-translucide */
                    color: #000;                             /* texte */
                    border-radius: 8px;
                    border: 1px solid rgba(255, 255, 255, 0.25);
                    selection-background-color: rgba(120, 150, 255, 0.5);
                    selection-color: #ffffff;
                }

                """
            )
            
            widget.academy_combo.clear()
            widget.academy_combo.addItems(self.data['AE'].unique().tolist())

            widget.scrollAreaWidgetContents.layout().insertWidget(0,self.modify_combo)

            widget.enregistrer.clicked.connect(lambda: self.modify_save(widget,widget.vice_group.layout(), deleteVar, dialog))
        
            content_layout = dialog.container.layout()
            content_layout.addWidget(widget)

            dialog.exec()
        except Exception as e:
            logging.error("Une erreur est survenue", exc_info=True)  # Enregistre l'erreur avec stack trace
    # Fonction pour collecter et afficher les données du centre selectionné
    def modify_data(self,widget, deleteVar:list=None):
        try:
            centre_data = self.data[self.data['Centre']==self.modify_combo.currentText()] # Données du centre selectionné
            centre_nb_line = centre_data.shape[0] # Nombre de lignes concernées

            academie = centre_data.iloc[0,0]
            centre = centre_data.iloc[0,1]
            examen = centre_data.iloc[0,-2]
            nb_salle = centre_data.iloc[0,-1]
            
            # Remplissage automatique des champs
            widget.academy_combo.setCurrentText(academie)
            widget.centre_line.setText(centre)
            widget.examen_combo.setCurrentText(examen)
            widget.salle_spin.setValue(int(nb_salle))

            widget.presi_prenom.setText(centre_data.iloc[0]['Prénom'])
            widget.presi_nom.setText(centre_data.iloc[0]['Nom'])
            widget.presi_matricule.setText(centre_data.iloc[0]['Matricule'])
            widget.presi_service.setText(centre_data.iloc[0]['Service'])
            widget.presi_categorie.setCurrentText(centre_data.iloc[0]['Catégorie'])
            widget.presi_poste.setText(centre_data.iloc[0]['Poste'])
            widget.presi_telephone.setText(centre_data.iloc[0]['Telephone'])

            vice_presi_layout = widget.vice_group.layout()  # Récupérer le layout actuel

            if vice_presi_layout:
                while vice_presi_layout.count():  # Supprimer tous les widgets du layout
                    child = vice_presi_layout.takeAt(0)
                    if child.widget():
                        child.widget().deleteLater()
            else:
                vice_presi_layout = QVBoxLayout()
                widget.vice_group.setLayout(vice_presi_layout)

            for i in range(1,centre_nb_line):
                
                vice_widget = Vice_president(f"prenom{i}",f"nom{i}",f"matricule{i}",f"service{i}",f"categorie{i}",f"poste{i}",f"telephone{i}",vice_presi_layout, deleteVar)
                
                vice_widget.prenom_line.setText(centre_data.iloc[i]['Prénom'])
                vice_widget.nom_line.setText(centre_data.iloc[i]['Nom'])
                vice_widget.matricule_line.setText(centre_data.iloc[i]['Matricule'])
                vice_widget.service_line.setText(centre_data.iloc[i]['Service'])
                vice_widget.categorie_combo.setCurrentText(centre_data.iloc[i]['Catégorie'])
                vice_widget.poste_line.setText(centre_data.iloc[i]['Poste'])
                vice_widget.telephone_line.setText(centre_data.iloc[i]['Telephone'])
                vice_widget.matricule_line.setReadOnly(True) # Rendre non modifiable le champ matricule

                vice_presi_layout.addWidget(vice_widget)

            Vice_president_add = QToolButton()
            Vice_president_add.setStyleSheet("background-color: rgba(37, 107, 182, 200); border-radius: 8px; border: 1px solid rgba(255,255,255,0.40)")
            Vice_president_add.setText("Ajouter")
            Vice_president_add.setToolTip("Ajouter un vice président")
            Vice_president_add.setFixedHeight(40)
            Vice_president_add.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
            Vice_president_add.setCursor(QCursor(Qt.CursorShape.PointingHandCursor))
            Vice_president_add.clicked.connect(lambda: self.vice_president_add_slot(vice_presi_layout))
            vice_presi_layout.addWidget(Vice_president_add)
            
        except Exception as e:
            logging.error("Une erreur est survenue", exc_info=True)  # Enregistre l'erreur avec stack trace
    def modify_save(self,widget,vice_layout, deleteVar:list=None, transparentBgDialog=None):
        loading = LoadingPage()
        loading.show()
        self.modifyThread = ModifyThread(widget,vice_layout, deleteVar, transparentBgDialog)
        self.modifyThread.process.connect(lambda message: self.statusBar().showMessage(message))
        self.modifyThread.finished.connect(lambda success: self.combo_change() if success else None)
        self.modifyThread.finished.connect(lambda success: loading.close() if success else None)
        self.modifyThread.finished.connect(lambda success: self.statusBar().clearMessage())
        self.modifyThread.finished.connect(lambda success: QMessageBox.information(self,"Succès","Modification apportée avec succès !") if success else QMessageBox.critical(self,"Erreur","Une erreur est survenue lors de l'enregistrement."))
        self.modifyThread.finished.connect(lambda: self.modifyThread.quit)
        self.modifyThread.finished.connect(lambda: self.modifyThread.wait)
        self.modifyThread.start()

class DataUploadThread(QThread):
    academieList = pyqtSignal(list)
    concatenateDf = pyqtSignal(pd.DataFrame)
    finished = pyqtSignal(bool)
    def __init__(self, file_path):
        super().__init__()
        self.file_path = file_path
    def run(self):
        try:
            new_file = pd.read_excel(self.file_path)
            new_file.fillna("N/A", inplace=True)
            new_file.sort_values(by=["AE","Centre"],inplace=True)
            conn = sqlite3.connect(os.path.join(folder, "database.db"))
            existed_data = pd.read_sql_query("select * from academy",conn)
            concatened_df = pd.concat([existed_data,new_file], ignore_index=True)
            concatened_df.sort_values(by=["AE","Centre"],inplace=True)
            self.concatenateDf.emit(concatened_df)
            new_file.to_sql(name="academy",con=conn, if_exists="append", index=False)
            academies_list = new_file['AE'].unique().tolist()
            self.academieList.emit(academies_list)
            conn.close()
            logging.info("Données importées avec succès !")
            self.finished.emit(True)
        except:
            logging.error("Une erreur est servenue lors de l'importation des données !", exc_info=True)


class TableFillThread(QThread):
    progress = pyqtSignal(int)           # %
    finished = pyqtSignal(list, list)    # rows, columns
    error = pyqtSignal(list)

    def __init__(self, data: pd.DataFrame):
        super().__init__()
        self.data = data
        self._running = True

    def stop(self):
        self._running = False

    def run(self):
        try:
            rows = []
            columns = list(self.data.columns)
            total = self.data.shape[0]

            for i in range(total):
                if not self._running:
                    return  # arrêt propre

                row_data = [
                    str(self.data.iat[i, j])
                    for j in range(self.data.shape[1])
                ]
                rows.append(row_data)

                percent = int((i + 1) / total * 100)
                self.progress.emit(percent)

            self.finished.emit(rows, columns)

        except Exception:
            columns = [
                "AE","Centre","Responsabilité","Prénom","Nom","Matricule",
                "Service","Catégorie","Poste","Telephone","Examen","Nb salle"
            ]
            self.error.emit(columns)

class ModifyThread(QThread):
    process = pyqtSignal(str)
    finished = pyqtSignal(bool)
    def __init__(self, widget,vice_layout, deleteVar:list=None, transparentBgDialog=None):
        super().__init__()
        self.widget = widget
        self.vice_layout = vice_layout
        self.deleteVar = deleteVar
        self.transparentBgDialog = transparentBgDialog
    def run(self):
        try:
            self.process.emit("Enregistrement en cours...")
            academy = self.widget.academy_combo.currentText()
            centre = self.widget.centre_line.text()
            examen = self.widget.examen_combo.currentText()
            nb_salle = self.widget.salle_spin.value()
            presi_prenom = self.widget.presi_prenom.text()
            presi_nom = self.widget.presi_nom.text()
            presi_matricule = self.widget.presi_matricule.text()
            presi_service = self.widget.presi_service.text()
            presi_categorie = self.widget.presi_categorie.currentText()
            presi_poste = self.widget.presi_poste.text()
            presi_telephone = self.widget.presi_telephone.text()
            president = ["Président",presi_prenom,presi_nom,presi_service,presi_categorie,
                presi_poste,presi_telephone,examen,nb_salle, academy,centre,presi_matricule]
            req = """
            UPDATE academy 
            SET Responsabilité=?,Prénom=?,Nom=?,Service=?,Catégorie=?,Poste=?,Telephone=?,Examen=?, "Nb salle"=? 
            WHERE AE=? AND Centre=? AND Matricule=?
            """
            req
            conn = sqlite3.connect(os.path.join(folder, "database.db"))
            cur = conn.cursor()
            cur.execute(req, president)
            vice_presi_count = self.vice_layout.count()-1
            for i in range(vice_presi_count):
                widget = self.vice_layout.itemAt(i).widget() # Récupérer le widget de chaque vice président
                prenom = widget.prenom_line.text()
                nom = widget.nom_line.text()
                matricule = widget.matricule_line.text()
                service = widget.service_line.text()
                categorie = widget.categorie_combo.currentText()
                poste = widget.poste_line.text()
                telephone = widget.telephone_line.text()
                data = [
                    f"{i+1}-Vice Président", prenom, nom, service, categorie, poste, telephone,
                    examen, nb_salle, academy, centre, matricule
                ]
                if widget.status == "new":
                    data = [
                        academy, centre, f"{i+1}-Vice Président", prenom, nom, matricule, service, categorie, poste, telephone, examen, nb_salle
                    ]
                    cur.execute("""INSERT INTO academy VALUES(?,?,?,?,?,?,?,?,?,?,?,?)""",data)
                else:
                    cur.execute(req, data)
            if self.deleteVar and not len(self.deleteVar) == 0:
                for matricule in self.deleteVar:
                    cur.execute("DELETE FROM academy WHERE Matricule=?", (matricule,))
                    logging.info(f"Suppréssion d'un vice président {matricule} réussie.")
            conn.commit()
            conn.close()
            self.deleteVar.clear()
            self.deleteVar.append(1)
            self.transparentBgDialog.close()
            self.finished.emit(True)
        except Exception as e:
            logging.error("Une erreur est survenue", exc_info=True)  # Enregistre l'erreur avec stack trace
            self.finished.emit(False)
class GlassOverlay(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setGeometry(parent.rect())
        self.setWindowFlags(Qt.FramelessWindowHint | Qt.Dialog)
        self.setAttribute(Qt.WA_TranslucentBackground)
        self.setAttribute(Qt.WA_DeleteOnClose)

        # Layout central pour le dialog
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setAlignment(Qt.AlignCenter)

        # Voile semi-transparent (optionnel, pour le contraste)
        veil = QWidget(self)
        veil.setStyleSheet("background-color: rgba(255,255,255,150); border-radius: 0px;")
        veil.setGeometry(self.rect())

        # Blur effect sur le parent
        blur_effect = QGraphicsBlurEffect()
        blur_effect.setBlurRadius(15)
        veil.setGraphicsEffect(blur_effect)

        # Dialog de chargement
        self.loading = LoadingPage()
        layout.addWidget(self.loading)
        self.loading.show()

class LoadingPage(QDialog):
    def __init__(self):
        super().__init__()
        self.setWindowFlags(Qt.FramelessWindowHint)
        self.setAttribute(Qt.WA_TranslucentBackground)
        self.setModal(True)

        layout = QVBoxLayout()
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(15)

        # GIF
        label = QLabel()
        label.setFixedSize(250, 250)
        label.setAlignment(Qt.AlignCenter)
        movie = QMovie(resource_path("icons/loading2.gif"))
        movie.setScaledSize(label.size())
        label.setMovie(movie)
        movie.start()
        layout.addWidget(label, alignment=Qt.AlignCenter)

        self.setLayout(layout)
        self.setStyleSheet("""
            background-color: rgba(255,255,255,150);
            border-radius: 12px;
        """)

class GenerateWordThread(QThread):
    progress = pyqtSignal(int)
    finish = pyqtSignal(bool)
    request_save_path = pyqtSignal()
    save_path_received = pyqtSignal(str)
    def __init__(self, parent=None):
        super().__init__(parent)
        self.parent = parent
        self.doc = None

    def run(self):
        conditions = {1:[1,2,3],2:[4,5,6,7],3:[8,9,10,11],4:[12,13,14,15],5:[16,17,18,19],6:[20,21,22,23],7:[24,25,26,27],
                          8:[28,29,30,31],9:[32,33,34,35], 10:[36,37,38,39,40]}
        try:
            con = sqlite3.connect(os.path.join(folder, "database.db"))
            main_df = pd.read_sql_query("select * from academy",con)

            dfs = query(self.parent.widget.academy_combo)
            academies = main_df['AE'].unique().tolist() # Avoir la liste des academies
            academies_dict = {} # Dictionnaire qui contiendra les noms d'academies comme clé et 0 comme valeur
            for academie in academies:
                academies_dict[academie] = 0
            doc = Document()
            style = doc.styles['Normal']
            style.paragraph_format.line_spacing = Pt(10)
            local_progress = 0
            for i,df in enumerate(dfs):
                if not df.empty:
                    if academies_dict[df.iloc[0,0]] == 0: # Si la valeur d'une academie est 0 donc elle n'est pas encore afficher
                        academiesLen = main_df["AE"].unique().tolist() # Obtenir le nombre d'académies
                        current_academie_index = academiesLen.index(df.iloc[0,0])# Obtenir l'index de l'académie actuelle
                        titre = doc.add_paragraph(style="Heading 1")
                        titre.add_run(f"{current_academie_index+1}/ ACADÉMIE D'ENSEIGNEMENT DE : {df.iloc[0,0]}").bold = True
                        titre.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        academies_dict[df.iloc[0,0]] += 1

                        # Extraction des superviseurs
                        superviseurs = main_df[main_df['Nb salle'].isnull()]
                        superviseurs = superviseurs[superviseurs['AE']==df.iloc[0,0]]

                        # Ajouter la section Superviseurs
                        doc.add_paragraph("SUPERVISEURS :", style="Heading 2")
                        # Création du tableau principal
                        table = doc.add_table(rows=superviseurs.shape[0]+1, cols=5)
                        table.style = "Table Grid"
                        # Remplir l'en-tête du tableau
                        hdr_cells = table.rows[0].cells
                        headers = ["N° d'ordre", "Prénoms", "Noms", "N° Mle", "Service"]
                        for h, text in enumerate(headers):
                            hdr_cells[h].text = text
                        for i in range(superviseurs.shape[0]):
                            superviseur_prenom = superviseurs.iloc[i]['Prénom']
                            superviseur_nom = superviseurs.iloc[i]['Nom']
                            superviseur_matricule = superviseurs.iloc[i]['Matricule']
                            superviseur_service = superviseurs.iloc[i]['Service']
                            row_cells = table.rows[i].cells
                            row_data = ["-", superviseur_prenom, superviseur_nom, superviseur_matricule, superviseur_service]
                            for i, text in enumerate(row_data):
                                row_cells[i].text = str(text)
                    doc.add_paragraph("", style=style)  # Ajouter un saut de ligne entre les centres
                    # Ajouter le titre du centre
                    centre = df.iloc[0,1]
                    academieCenterLen = main_df[main_df['AE']==df.iloc[0,0]].Centre.unique().tolist() # Obtenir le nombre de centres de cette académie
                    centre_index = academieCenterLen.index(centre) + 1 # Obtenir l'index du centre actuel
                    salle = int(df.iloc[0,-1])
                    titre_a = doc.add_paragraph()
                    titre_a.add_run(f"{centre_index}- CENTRE DU {centre}").bold = True
                    titre_a.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    titre_a.add_run(" " * 50 + f"{str(salle)}").bold = True  # Espacer le numéro du centre
                    # Ajouter la section RESPONSABLE
                    doc.add_paragraph("RESPONSABLE :", style="Heading 2")
                    # Création du tableau principal
                    table = doc.add_table(rows=2, cols=5)
                    table.style = "Table Grid"
                    # Remplir l'en-tête du tableau
                    hdr_cells = table.rows[0].cells
                    headers = ["N° d'ordre", "Prénoms", "Noms", "N° Mle", "Service"]
                    for h, text in enumerate(headers):
                        hdr_cells[h].text = text
                    # Ajouter le responsable
                    presidents_liste = presidents(self.parent.widget.academy_combo)
                    current_president = presidents_liste[presidents_liste['Centre']==centre]
                    # Si aucune ligne trouvée
                    if current_president.empty:
                        president_prenom = ""
                        president_nom = ""
                        president_matricule = ""
                        president_service = ""
                    else:
                        ligne = current_president.iloc[0]  # On prend la première ligne trouvée

                        president_prenom = ligne.get('Prénom', "")
                        president_nom = ligne.get('Nom', "")
                        president_matricule = ligne.get('Matricule', "")
                        president_service = ligne.get('Service', "")
                    row_cells = table.rows[1].cells
                    numerotation = 1 # Numérotation des lignes de noms
                    row_data = [str(numerotation), president_prenom, president_nom, president_matricule, president_service]
                    for i, text in enumerate(row_data):
                        row_cells[i].text = str(text)

                    # Ajouter la section RESPONSABLES ADJOINTS
                    doc.add_paragraph("RESPONSABLES ADJOINTS :", style="Heading 2")

                    for condition, values in conditions.items():
                        if salle in values:
                            nb_responsable_adjoint = condition
                            break
                    # Création du second tableau
                    table2 = doc.add_table(rows=nb_responsable_adjoint+1, cols=5)
                    table2.style = "Table Grid"

                    # En-tête du tableau
                    hdr_cells2 = table2.rows[0].cells
                    for i, text in enumerate(headers):
                        hdr_cells2[i].text = text

                    # Ajouter les responsables adjoints
                    rows_data = []
                    for i in range(df.shape[0]):
                        numerotation += 1
                        rows_data.append([str(numerotation),df.iloc[i]["Prénom"],df.iloc[i].Nom, df.iloc[i].Matricule, df.iloc[i].Service])
                    numerotation = 1
                    for i, row_data in enumerate(rows_data):
                        try:
                            row_cells = table2.rows[i+1].cells
                            for j, text in enumerate(row_data):
                                row_cells[j].text = text
                        except Exception as e:
                            logging.error("Une erreur est survenue", exc_info=True)  # Enregistre l'erreur avec stack trace
                local_progress += 1
                self.progress.emit(int((local_progress*100)/len(dfs)))
            self.doc = doc  # garder le document en mémoire
            self.request_save_path.emit()

            if file_path:  # Vérifie si un chemin a été sélectionné
                if not file_path.endswith(".docx"):  # Ajoute l'extension si nécessaire
                    file_path += ".docx"
                doc.save(file_path)
                self.finish.emit(True)
                return
            self.finish.emit(False)
        except Exception as e:
            logging.error("Une erreur est survenue", exc_info=True)  # Enregistre l'erreur avec stack trace

class Vice_president(QFrame):
    """
    Le QFrame représentant un vice président dans la boîte d'ajout/modification.
    Params:
    - prenom: str
    - nom: str
    - matricule: str
    - service: str
    - categorie: str
    - poste: str
    - telephone: str
    - vice_group_layout: QVBoxLayout (layout parent dans lequel sera ajouté les frames des vices présidents)
    - deleteVar: La variable qui contiendra le numéro du vice président (utile pour la suppression)
    """
    def __init__(self,prenom,nom,matricule,service,categorie,poste, telephone, vice_group_layout:QVBoxLayout, deleteVar=None, status:str=None):
        super().__init__()
        self.prenom = prenom
        self.nom = nom
        self.matricule = matricule
        self.service = service
        self.categorie = categorie
        self.poste = poste
        self.telephone = telephone
        self.vice_group_layout = vice_group_layout
        self.deleteVar = deleteVar
        self.status = status

        layout = QVBoxLayout()
        self.prenom_line = Line(self.prenom,"Prénom")
        self.nom_line = Line(self.nom, "Nom")
        self.matricule_line = Line(self.matricule, "Matricule")
        self.service_line = Line(self.service, "Service")
        categorie_frame = QFrame()
        categorie_frame.setStyleSheet("""
            QFrame {
                background: rgba(255, 255, 255, 0.20);       /* fond translucide */
                border: 1px solid rgba(255, 255, 255, 0.25);
                border-radius: 8px;
            }

            /* Hover : léger éclaircissement */
            QFrame:hover {
                background: rgba(255, 255, 255, 0.30);
                border: 1px solid rgba(120, 150, 255, 0.6);
            }

            /* Focus (si le frame reçoit le focus) */
            QFrame:focus {
                background: rgba(255, 255, 255, 0.45);
                border: 1px solid rgba(120, 150, 255, 0.8);
            }
        """)
        categorie_layout = QHBoxLayout()
        categorie_layout.setContentsMargins(0,0,0,0)
        categorie_label = QLabel(text="Categorie")
        categorie_label.setStyleSheet("background:transparent; border:none; color: #000;")
        self.categorie_combo = QComboBox()
        self.categorie_combo.setStyleSheet("background: transparent; border: none; color: #000;")
        self.categorie_combo.setMinimumHeight(50)
        self.categorie_combo.setStyleSheet("border:0")
        self.categorie_combo.setObjectName(self.categorie)
        self.categorie_combo.addItems(["A","B","C","D","E"])
        categorie_layout.addWidget(categorie_label,0)
        categorie_layout.addWidget(self.categorie_combo,1)
        categorie_frame.setLayout(categorie_layout)
        self.poste_line = Line(self.poste, "Poste")
        self.telephone_line = Line(self.telephone, "Telephone")

        remove_btn = QToolButton()  # Button permettant de supprimer un vice président
        remove_btn.setFixedSize(150,35)
        remove_btn.setStyleSheet("border-radius:10px; background-color:rgb(195,69,69); border:0")
        remove_btn.setText("Supprimer")
        remove_btn.setCursor(QCursor(Qt.CursorShape.PointingHandCursor))
        remove_btn.clicked.connect(self.remove_slot)
        rm_btn_layout = QHBoxLayout()
        rm_btn_layout.setAlignment(Qt.AlignCenter)
        rm_btn_layout.addWidget(remove_btn)

        layout.addWidget(self.prenom_line)
        layout.addWidget(self.nom_line)
        layout.addWidget(self.matricule_line)
        layout.addWidget(self.service_line)
        layout.addWidget(categorie_frame)
        layout.addWidget(self.poste_line)
        layout.addWidget(self.telephone_line)
        layout.addLayout(rm_btn_layout)

        self.setLayout(layout)

        self.setStyleSheet("""
            border: 1px solid gray;
            border-radius: 10px;
                           """)
    # Slot du button de suppression de vice président
    def remove_slot(self):
        if self.deleteVar:
            self.deleteVar.append(self.matricule_line.text())
        self.vice_group_layout.removeWidget(self)
        
class Line(QLineEdit):
    def __init__(self,name,placeholder):
        super().__init__()
        self.name = name
        self.placeholder = placeholder
        s = """
            QLineEdit {
                background: rgba(255, 255, 255, 0.20);   /* fond translucide */
                color: #000;                           /* texte clair */
                border: 1px solid rgba(255, 255, 255, 100);
                border-radius: 8px;
                padding: 6px 10px;                        /* padding interne */
                selection-background-color: rgba(120, 150, 255, 0.5);
                selection-color: #ffffff;
                font-size: 15px;
            }

            QLineEdit:hover {
                background: rgba(255, 255, 255, 0.30);
                border: 1px solid rgba(120, 150, 255, 0.6);
            }

            QLineEdit:focus {
                background: rgba(255, 255, 255, 0.5);
                border: 1px solid rgba(120, 150, 255, 0.8);
            }
        """
        self.setStyleSheet(s)
        self.setObjectName(name)
        self.setMinimumHeight(50)
        self.setPlaceholderText(self.placeholder)

def has_decimal_part(number):
    decimal_part, _ = math.modf(number)
    return decimal_part != 0

def query(academie_combo:QComboBox):
    conn = sqlite3.connect(os.path.join(folder, "database.db"))
    conditions = {1:[1,2,3],2:[4,5,6,7],3:[8,9,10,11],4:[12,13,14,15],5:[16,17,18,19],6:[20,21,22,23],7:[24,25,26,27],
                          8:[28,29,30,31],9:[32,33,34,35], 10:[36,37,38,39,40]} # Conditions de selection du nombre de vice présidents
    if academie_combo.currentText() == "Tout":
        df = pd.read_sql_query("select * from academy",conn)
        df.sort_values(by=["AE","Centre"],inplace=True)
        centres = df.Centre.unique() # Récuperer les valeurs uniques des centres
        if not df[df.Centre == centres[0]].iloc[0,-1]:
            centres = np.delete(centres, 0) # Supprimer les lignes ayant 0 salle (vise à éliminer les lignes des superviseurs)
        df_list = []
        for c in centres:
            centre_df = df[df.Centre == c] # Avoir uniquement un dataframe pour chaque centre
            nb_salle = int(df[df.Centre == c].iloc[0,-1])  # Avoir le nombre salle en prennant la valeur de la colonne 'nb_salle' de la ligne président
            for k,v in conditions.items():
                if int(nb_salle) in v:
                    nb_vice_president = k
            try:
                df_list.append(centre_df[1:nb_vice_president+1])
            except IndexError:
                df_list.append(centre_df[1:])
    else:
        df = pd.read_sql_query(f"select * from academy WHERE AE='{academie_combo.currentText()}'",conn)
        df.sort_values(by=["AE","Centre"],inplace=True)
        centres = df.Centre.unique() # Récuperer les valeurs uniques des centres
        if not df[df.Centre == centres[0]].iloc[0,-1]:
            centres = np.delete(centres, 0) # Supprimer les lignes ayant 0 salle (vise à éliminer les lignes des superviseurs)
        df_list = []
        for c in centres:
            centre_df = df[df.Centre == c] # Avoir uniquement un dataframe pour chaque centre
            nb_salle = int(df[df.Centre == c].iloc[0,-1]) # Avoir le nombre salle en prennant la valeur de la colonne 'nb_salle' de la ligne président
            for k,v in conditions.items():
                if int(nb_salle) in v:
                    nb_vice_president = k
            try:
                df_list.append(centre_df[1:nb_vice_president+1])
            except IndexError:
                df_list.append(centre_df[1:])
    return  df_list# Séléctionner les vice présidents nécésairent

def presidents(academie_combo:QComboBox):
    if academie_combo.currentText() == "Tout":
        conn = sqlite3.connect(os.path.join(folder, "database.db"))
        df = pd.read_sql_query("select * from academy",conn)
        president = df[df['Responsabilité'].isin(["1-Président", "1-président", "Président", "président", "President", "president"])].copy()
        president.fillna("N/A",inplace=True)
        president.reset_index()
        return president
    else:
        conn = sqlite3.connect(os.path.join(folder, "database.db"))
        df = pd.read_sql_query(f"select * from academy WHERE AE='{academie_combo.currentText()}'",conn)
        president = df[df['Responsabilité'].isin(["1-Président", "1-président", "Président", "président", "President", "president"])].copy()
        president.fillna("N/A",inplace=True)
        president.reset_index()
        return president

from PyQt5.QtCore import QPoint
from PyQt5.QtGui import QGuiApplication
import numpy as np

class FrostedDialog(QDialog):
    def __init__(self, parent=None, w=None, blur_radius=18):
        super().__init__(parent)
        self.w = w
        self.blur_radius = blur_radius

        self.setFixedSize(635,800)
        self.setContentsMargins(0,0,0,0)
        # Fenêtre frameless + fond transparent (on verra le QLabel flouté)
        self.setWindowFlags(Qt.FramelessWindowHint | Qt.Dialog)
        self.setAttribute(Qt.WA_TranslucentBackground)

        # -------------- BACKGROUND LABEL (will hold the captured background pixmap) --------------
        self.bg_label = QLabel(self)
        self.bg_label.setScaledContents(True)  # le pixmap s'étirera pour remplir
        self.bg_label.setGeometry(self.rect())

        blur = QGraphicsBlurEffect(self.bg_label)
        blur.setBlurRadius(self.blur_radius)
        self.bg_label.setGraphicsEffect(blur)

        # -------------- CONTAINER (widgets nets) --------------
        self.container = QWidget(self)
        self.container.setContentsMargins(0,0,0,0)
        self.container.setObjectName("container")
        # style : semi-translucide pour effet glass mais contenu net
        self.container.setStyleSheet("""
            #container {
                background: rgba(20, 20, 20, 0.0);
                border-radius: 14px;
            }
        """)
        # padding interne
        self.container.setGeometry(self.rect().adjusted(12, 12, -12, -12))

        # Exemple d'UI à l'intérieur
        cl = QVBoxLayout(self.container)
        cl.addWidget(self.w)

    # ---- showEvent : première capture ----
    def showEvent(self, event):
        super().showEvent(event)
        self.updateBackground()

    # ---- moveEvent / resizeEvent : mettre à jour le fond quand le dialog bouge ----
    def moveEvent(self, event):
        super().moveEvent(event)
        self.updateBackground()

    def resizeEvent(self, event):
        super().resizeEvent(event)
        # redimensionner bg_label et container
        self.bg_label.setGeometry(self.rect())
        self.container.setGeometry(self.rect().adjusted(12, 12, -12, -12))
        self.updateBackground()

    # ---- Fonction qui capture l'écran derrière le dialog et la met en background label ----
    def updateBackground(self):
        # position globale du coin supérieur gauche du dialog
        top_left_global = self.mapToGlobal(QPoint(0, 0))
        x, y = top_left_global.x(), top_left_global.y()
        w, h = self.width(), self.height()

        # Trouver l'écran à cette position (gère multi-écrans)
        screen = QGuiApplication.screenAt(top_left_global)
        if screen is None:
            screen = QGuiApplication.primaryScreen()

        # Capture entière de l'écran en question
        full_pix = screen.grabWindow(0)

        # Si l'application est en multi-écran, il est possible que full_pix corresponde
        # à l'écran cible mais les coordonnées globales doivent être transformées.
        # On calcule la géométrie de l'écran pour connaître son offset dans l'espace global.
        screen_geo = screen.geometry()
        # offsets relatifs au pixmap
        rel_x = x - screen_geo.x()
        rel_y = y - screen_geo.y()

        # éviter les découpes hors limites
        rel_x = max(0, min(rel_x, full_pix.width() - 1))
        rel_y = max(0, min(rel_y, full_pix.height() - 1))
        w = min(w, full_pix.width() - rel_x)
        h = min(h, full_pix.height() - rel_y)
        if w <= 0 or h <= 0:
            return

        # Crop : zone correspondant à la position du dialog
        cropped = full_pix.copy(rel_x, rel_y, w, h)

        # Optionnel : mettre à l'échelle si on veut l'étirer pour correspondre à la taille du label
        # Ici bg_label est en ScaledContents=True, donc on peut fournir cropped tel quel.
        self.bg_label.setPixmap(cropped)



# Lancer l'application
app = QApplication([])
window = MyApp()
window.show()
app.exec_()

"""
28 - 31 = 8
"""